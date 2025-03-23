"""
Word文档处理器模块
"""
import os
from typing import Dict, Any, Union, List
import pandas as pd
import logging
import io

# 配置日志
logger = logging.getLogger(__name__)

# 导入基类
from app.services.file_service import BaseFileProcessor

class WordProcessor(BaseFileProcessor):
    """Word文档处理器"""
    name = "word_processor"
    description = "处理Word文档文件"
    supported_extensions = ['docx', 'doc', 'odt']
    supported_mime_types = [
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'application/msword',
        'application/vnd.oasis.opendocument.text'
    ]
    priority = 20
    
    @classmethod
    def extract_text(cls, file_path: str) -> str:
        """从Word文档中提取文本内容"""
        try:
            # 尝试使用python-docx
            import docx
            
            doc = docx.Document(file_path)
            full_text = []
            
            # 提取段落文本
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    full_text.append(" | ".join(row_text))
            
            return "\n".join(full_text)
                
        except ImportError:
            logger.error("python-docx未安装，无法提取Word文本")
            
            # 尝试使用pywin32（仅Windows）
            try:
                import win32com.client
                import os
                
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                
                abs_path = os.path.abspath(file_path)
                doc = word.Documents.Open(abs_path)
                text = doc.Content.Text
                
                doc.Close()
                word.Quit()
                
                return text
            except ImportError:
                logger.error("win32com未安装，无法提取Word文本")
                return "无法提取Word内容：需要安装python-docx或win32com库"
            except Exception as e:
                logger.error(f"使用win32com提取Word文本失败: {str(e)}")
                return f"Word解析错误: {str(e)}"
        except Exception as e:
            logger.error(f"Word文本提取错误: {str(e)}", exc_info=True)
            return f"Word解析错误: {str(e)}"
    
    @classmethod
    def extract_metadata(cls, file_path: str) -> Dict[str, Any]:
        """提取Word文档的元数据"""
        metadata = super().extract_metadata(file_path)
        
        try:
            import docx
            doc = docx.Document(file_path)
            
            # 添加文档特有元数据
            core_properties = doc.core_properties
            
            # 提取核心属性
            doc_info = {}
            properties = [
                'author', 'category', 'comments', 'content_status', 
                'created', 'identifier', 'keywords', 'language', 
                'last_modified_by', 'last_printed', 'modified', 
                'revision', 'subject', 'title', 'version'
            ]
            
            for prop in properties:
                if hasattr(core_properties, prop):
                    value = getattr(core_properties, prop)
                    if value is not None:
                        doc_info[prop] = str(value)
            
            metadata["document_info"] = doc_info
            
            # 提取文档统计信息
            metadata.update({
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            })
                
        except Exception as e:
            logger.error(f"提取Word元数据时出错: {str(e)}")
            
        return metadata
    
    @classmethod
    def extract_structured_data(cls, file_path: str) -> Union[pd.DataFrame, Dict, List, None]:
        """尝试从Word提取表格数据"""
        try:
            import docx
            doc = docx.Document(file_path)
            
            # 如果文档中有表格
            if doc.tables:
                all_tables = []
                
                for i, table in enumerate(doc.tables):
                    # 提取表格数据
                    data = []
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            row_data.append(cell.text)
                        data.append(row_data)
                    
                    # 如果表格非空
                    if data and len(data) > 0:
                        # 使用第一行作为列名
                        headers = data[0]
                        table_data = data[1:]
                        
                        # 创建DataFrame
                        df = pd.DataFrame(table_data, columns=headers)
                        all_tables.append(df)
                
                # 返回结果
                if len(all_tables) == 1:
                    return all_tables[0]
                elif len(all_tables) > 1:
                    result = {}
                    for i, table in enumerate(all_tables):
                        result[f"table_{i+1}"] = table
                    return result
                    
        except Exception as e:
            logger.error(f"提取Word结构化数据失败: {str(e)}", exc_info=True)
        
        return None 